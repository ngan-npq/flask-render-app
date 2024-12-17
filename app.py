import pandas as pd
import os
import requests
from flask import Flask, render_template, request, send_file
from datetime import datetime, timedelta
from io import BytesIO, StringIO
from docx import Document  # For Word document processing
from docx2pdf import convert
import zipfile

app = Flask(__name__)

# Google Sheet links for separate spreadsheets
SHEET_LINKS = {
    "license_data": "https://docs.google.com/spreadsheets/d/1FVqKScY2fS1OQ9KjnU_PZHnBXkNQP9ZAm1kIPOOAFJE/export?format=csv",
    "ro": "https://docs.google.com/spreadsheets/d/1j5gr2FA2sIAESdmzC0vuRbVS3cdBH2QzjoHDv4XXNn8/export?format=csv",
    "prof_coll": "https://docs.google.com/spreadsheets/d/17MB2mBm0tvHkRVMh0IKJOZLblRrHo69Vuf5YbS0dY4A/export?format=csv",
}

LOG_FILE = "output/log_records.csv"
ROMAN_MONTHS = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII']

# Helper function to fetch CSV data
def fetch_csv_data(link):
    response = requests.get(link)
    return pd.read_csv(StringIO(response.text))

# Helper function to append logs
def append_to_log(record):
    os.makedirs("output", exist_ok=True)
    if not os.path.exists(LOG_FILE):
        pd.DataFrame(columns=record.keys()).to_csv(LOG_FILE, index=False)
    log_data = pd.read_csv(LOG_FILE)
    log_data = pd.concat([log_data, pd.DataFrame([record])], ignore_index=True)
    log_data.to_csv(LOG_FILE, index=False)

@app.route('/')
def index():
    # Fetch license plate data
    license_data = fetch_csv_data(SHEET_LINKS["license_data"])
    license_plates = license_data['license_plate'].dropna().tolist()

    # Fetch Collection PIC names
    internal_pics = fetch_csv_data(SHEET_LINKS["ro"]).iloc[:, 0].dropna().tolist()  # First column
    external_pics = fetch_csv_data(SHEET_LINKS["prof_coll"]).iloc[:, 0].dropna().tolist()  # First column

    return render_template(
        "form.html",
        license_plates=license_plates,
        internal_pics=internal_pics,
        external_pics=external_pics
    )

def replace_placeholders(template, replacements):
    """
    Replaces placeholders in Word document paragraphs and tables,
    combining runs to handle placeholders spanning across multiple runs.
    """
    # Replace placeholders in paragraphs
    for paragraph in template.paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)  # Combine all runs
        for placeholder, replacement in replacements.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(replacement))
                # Update runs with the replaced text
                for run in paragraph.runs:
                    run.text = ""
                paragraph.runs[0].text = full_text

    # Replace placeholders in tables
    for table in template.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders(cell, replacements)  # Recursive call for table cells

@app.route('/generate_pdf', methods=['POST'])
def generate_pdf():
    # Collect form data
    license_plate = request.form['license_plate']
    internal_external = request.form['internal_external']
    collection_pic = request.form['collection_pic']
    user_name = request.form['user_name']

    # Fetch license data
    license_data = fetch_csv_data(SHEET_LINKS["license_data"])
    vehicle_row = license_data[license_data['license_plate'] == license_plate].iloc[0]

    # Fetch PIC data for alamat
    if internal_external == "Internal":
        pic_data = fetch_csv_data(SHEET_LINKS["ro"])
    else:
        pic_data = fetch_csv_data(SHEET_LINKS["prof_coll"])
    pic_row = pic_data[pic_data.iloc[:, 0] == collection_pic].iloc[0]
    alamat = pic_row.iloc[1]  # Second column is 'alamat'

    # Download Fiducia Certificate
    fiducia_link = vehicle_row.get('fiducia_certificate', None)
    fiducia_doc_path = None
    if fiducia_link:
        response = requests.get(fiducia_link)
        if response.status_code == 200:
            os.makedirs("output", exist_ok=True)
            fiducia_filename = f"output/{license_plate}_fiducia_certificate.pdf"
            with open(fiducia_filename, "wb") as f:
                f.write(response.content)
            fiducia_doc_path = fiducia_filename

    # Generate SK Auto Number
    current_month = datetime.now().strftime("%Y-%m")
    if os.path.exists(LOG_FILE):
        log_data = pd.read_csv(LOG_FILE)
        log_data['document_date'] = pd.to_datetime(log_data['document_date'], errors='coerce')
        log_data = log_data.dropna(subset=['document_date'])
        same_month_logs = log_data[log_data['document_date'].dt.strftime("%Y-%m") == current_month]
        sequence_number = len(same_month_logs) + 1  # Increment based on existing count
    else:
        sequence_number = 1

    sk_autogenerate_number = f"{sequence_number}/DIV.COLL-PST/DEFI/{ROMAN_MONTHS[datetime.now().month-1]}/{datetime.now().year}"
    document_date = datetime.now().strftime("%d %B %Y")
    sk_expiry_date = (datetime.now() + timedelta(days=7)).strftime("%d %B %Y")

    # Calculate total_amount dynamically
    principal = pd.to_numeric(vehicle_row.get('principal', 0), errors='coerce') or 0
    interest = pd.to_numeric(vehicle_row.get('interest', 0), errors='coerce') or 0
    penalty = pd.to_numeric(vehicle_row.get('penalty', 0), errors='coerce') or 0
    default_fee = pd.to_numeric(vehicle_row.get('default_fee', 0), errors='coerce') or 0
    total_amount = principal + interest + penalty + default_fee

    # Fill Word template
    template = Document("files/SK Template.docx")
    replacements = {
        "{{license_plate}}": vehicle_row['license_plate'],
        "{{nama_collection}}": collection_pic,
        "{{alamat}}": alamat,
        "{{arm}}": user_name,
        "{{owner_name}}": vehicle_row['owner_name'],
        "{{dealer_name}}": vehicle_row['dealer_name'],
        "{{dealer_address}}": vehicle_row['dealer_address'],
        "{{brand}}": vehicle_row['brand'],
        "{{model}}": vehicle_row['model'],
        "{{engine_number}}": vehicle_row['engine_number'],
        "{{chassis_number}}": vehicle_row['chassis_number'],
        "{{car_year}}": vehicle_row['car_year'],
        "{{car_color}}": vehicle_row['car_color'],
        "{{dpd}}": vehicle_row['dpd'],
        "{{principal}}": principal,
        "{{interest}}": interest,
        "{{penalty}}": penalty,
        "{{default_fee}}": default_fee,
        "{{total_amount}}": total_amount,
        "{{document_date}}": document_date,
        "{{sk_expiry_date}}": sk_expiry_date,
        "{{sk_autogenerate_number}}": sk_autogenerate_number,
    }

    replace_placeholders(template, replacements)

    # Save the filled Word document temporarily
    word_path = "output/temp_filled_template.docx"
    os.makedirs("output", exist_ok=True)
    template.save(word_path)

    # Generate a dynamic PDF filename
    sanitized_expiry_date = sk_expiry_date.replace(" ", "_")  # Replace spaces with underscores
    pdf_filename = f"{license_plate}_{vehicle_row['dealer_name']}_{sanitized_expiry_date}.pdf"
    pdf_path = f"output/{pdf_filename}"

    # Convert the Word document to PDF
    convert(word_path, pdf_path)

    # Log the record
    log_record = {
        "sk_autogenerate_number": sk_autogenerate_number,
        "document_date": document_date,
        "sk_expiry_date": sk_expiry_date,
        "license_plate": vehicle_row['license_plate'],
        "owner_name": vehicle_row['owner_name'],
        "dealer_name": vehicle_row['dealer_name'],
        "dealer_address": vehicle_row['dealer_address'],
        "brand": vehicle_row['brand'],
        "model": vehicle_row['model'],
        "engine_number": vehicle_row['engine_number'],
        "chassis_number": vehicle_row['chassis_number'],
        "collection_pic": collection_pic,
        "alamat": alamat,
        "principal": principal,
        "interest": interest,
        "penalty": penalty,
        "default_fee": default_fee,
        "total_amount": total_amount,
        "arm": user_name,
        "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    append_to_log(log_record)

    # Paths to the files
    files_to_zip = [pdf_path]  # Include the generated PDF
    if fiducia_doc_path:  # Add fiducia_certificate if it exists
        files_to_zip.append(fiducia_doc_path)
    
    # Create a ZIP file
    zip_filename = f"{license_plate}_documents.zip"
    zip_path = f"output/{zip_filename}"

    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file in files_to_zip:
            zipf.write(file, os.path.basename(file))  # Add files to the ZIP archive
    
    # Send the ZIP file as the response
    return send_file(zip_path, as_attachment=True, download_name=zip_filename)

if __name__ == "__main__":
    app.run(debug=True)